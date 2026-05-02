"""
Document text extractor — converts uploaded files to clean text via Bedrock.

Supported formats:
  PDF           → native Bedrock Claude document message (base64)
  DOCX          → python-docx → raw text → Bedrock structures it
  TXT           → raw text → Bedrock cleans/structures it
  CSV           → csv stdlib → markdown table → Bedrock
  Excel (.xlsx) → openpyxl → markdown table → Bedrock
"""

import csv
import io
from pathlib import Path

from app.core.bedrock import extract_document as _bedrock_extract


def extract(file_bytes: bytes, filename: str) -> str:
    """Extract text from a file using the appropriate strategy.

    Args:
        file_bytes: Raw file bytes.
        filename:   Original filename (used to determine format).

    Returns:
        Cleaned, structured text string.

    Raises:
        ValueError: If the file format is not supported.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(file_bytes, filename)
    elif ext == ".docx":
        return _extract_docx(file_bytes, filename)
    elif ext == ".txt":
        return _extract_txt(file_bytes, filename)
    elif ext == ".csv":
        return _extract_csv(file_bytes, filename)
    elif ext in (".xlsx", ".xls"):
        return _extract_excel(file_bytes, filename)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


# ── Format-specific extractors ────────────────────────────────────────────────

def _extract_pdf(file_bytes: bytes, filename: str) -> str:
    """Send PDF to Bedrock Claude as a native document message."""
    return _bedrock_extract(file_bytes, "application/pdf", filename)


def _extract_docx(file_bytes: bytes, filename: str) -> str:
    """Extract text from DOCX using python-docx, then clean with Bedrock."""
    from docx import Document
    doc   = Document(io.BytesIO(file_bytes))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    raw_text = "\n".join(lines)
    return _bedrock_extract(
        raw_text.encode("utf-8"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename,
    )


def _extract_txt(file_bytes: bytes, filename: str) -> str:
    """Clean and structure plain text with Bedrock."""
    return _bedrock_extract(file_bytes, "text/plain", filename)


def _extract_csv(file_bytes: bytes, filename: str) -> str:
    """Convert CSV to markdown table, then summarise with Bedrock."""
    text      = file_bytes.decode("utf-8", errors="ignore")
    reader    = csv.reader(io.StringIO(text))
    rows      = list(reader)
    if not rows:
        return ""
    # Build markdown table (max 200 rows to keep prompt manageable)
    header    = rows[0]
    md_header = "| " + " | ".join(header) + " |"
    md_sep    = "| " + " | ".join(["---"] * len(header)) + " |"
    md_rows   = [
        "| " + " | ".join(str(c) for c in row) + " |"
        for row in rows[1:201]
    ]
    markdown  = "\n".join([md_header, md_sep] + md_rows)
    if len(rows) > 201:
        markdown += f"\n\n_(showing first 200 of {len(rows)-1} rows)_"
    return _bedrock_extract(markdown.encode("utf-8"), "text/csv", filename)


def _extract_excel(file_bytes: bytes, filename: str) -> str:
    """Convert Excel to markdown tables (one per sheet), then summarise with Bedrock."""
    import openpyxl
    wb      = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    parts   = []
    for sheet_name in wb.sheetnames:
        ws    = wb[sheet_name]
        rows  = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        parts.append(f"## Sheet: {sheet_name}")
        header  = [str(c) if c is not None else "" for c in rows[0]]
        md_head = "| " + " | ".join(header) + " |"
        md_sep  = "| " + " | ".join(["---"] * len(header)) + " |"
        md_rows = [
            "| " + " | ".join(str(c) if c is not None else "" for c in row) + " |"
            for row in rows[1:201]
        ]
        parts.append("\n".join([md_head, md_sep] + md_rows))
        if len(rows) > 201:
            parts.append(f"_(showing first 200 of {len(rows)-1} rows)_")
    wb.close()
    combined = "\n\n".join(parts)
    return _bedrock_extract(
        combined.encode("utf-8"),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename,
    )


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".xlsx", ".xls"}
SUPPORTED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "text/csv",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
}
